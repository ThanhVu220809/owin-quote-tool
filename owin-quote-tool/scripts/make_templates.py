# -*- coding: utf-8 -*-
"""
Sinh 2 file template Word cho docxtemplater (TASK 4.4).
Tag docxtemplater:
  - vòng lặp dòng: {#items} ... {/items} đặt trong CÙNG một hàng bảng → lặp hàng.
  - ảnh (Format 2): {%image} (image-module).
  - placeholder ngoài loop: {ten_kh} {dia_chi} {sdt} {email} {tong_tien} {tam_ung} {con_lai}.
Mỗi ô = 1 run để tag không bị Word cắt giữa chừng.
KHÔNG merge ô dọc (BR docx).
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "src", "assets", "templates")
os.makedirs(OUT_DIR, exist_ok=True)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)


def add_customer_block(doc):
    doc.add_paragraph().add_run("Khách hàng: {ten_kh}")
    doc.add_paragraph().add_run("Địa chỉ: {dia_chi}")
    doc.add_paragraph().add_run("SĐT: {sdt}    Email: {email}")


def add_totals_block(doc):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.add_run("Tổng cộng: {tong_tien}").bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("Tạm ứng: {tam_ung}")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("Còn lại: {con_lai}").bold = True


def build_format1():
    doc = Document()
    add_title(doc, "BÁO GIÁ CÔNG TRÌNH")
    add_customer_block(doc)
    doc.add_paragraph()

    headers = ["STT", "Mã", "Mô tả", "ĐVT", "Rộng", "Cao", "SL", "KL", "Đơn giá", "Thành tiền"]
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)

    # Hàng dữ liệu duy nhất, bọc {#items} ... {/items}
    row = table.rows[1].cells
    set_cell(row[0], "{#items}{stt}")
    set_cell(row[1], "{ma}")
    set_cell(row[2], "{mo_ta}")
    set_cell(row[3], "{dvt}")
    set_cell(row[4], "{rong}")
    set_cell(row[5], "{cao}")
    set_cell(row[6], "{sl}")
    set_cell(row[7], "{khoi_luong}")
    set_cell(row[8], "{don_gia}")
    set_cell(row[9], "{thanh_tien}{/items}")

    add_totals_block(doc)
    out = os.path.join(OUT_DIR, "Template_Bao_Gia.docx")
    doc.save(out)
    print("saved", out)


def build_format2():
    doc = Document()
    add_title(doc, "BẢNG GIÁ HOÀN THIỆN NHÔM OWIN")
    add_customer_block(doc)
    doc.add_paragraph()

    headers = ["STT", "Ảnh", "Mã", "Mô tả", "Kích thước", "ĐVT", "SL", "Đơn giá", "Thành tiền"]
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)

    row = table.rows[1].cells
    set_cell(row[0], "{#items}{stt}")
    set_cell(row[1], "{%image}")  # image-module, trong loop
    set_cell(row[2], "{ma}")
    set_cell(row[3], "{mo_ta}")
    set_cell(row[4], "{kich_thuoc}")
    set_cell(row[5], "{dvt}")
    set_cell(row[6], "{sl}")
    set_cell(row[7], "{don_gia}")
    set_cell(row[8], "{thanh_tien}{/items}")

    add_totals_block(doc)
    out = os.path.join(OUT_DIR, "Template_Bang_Gia.docx")
    doc.save(out)
    print("saved", out)


if __name__ == "__main__":
    build_format1()
    build_format2()
